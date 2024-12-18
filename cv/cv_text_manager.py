from docx import Document
from docx.shared import Pt, RGBColor

# Function to parse .docx file and extract text # Old function likely to be replaced
def extract_docx_text(file_path, Fsg):
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        Fsg.popup(f"Error reading file: {e}")
        return ""

def copy_docx_with_formatting(input_file, output_file):
    # Load the original document
    original_doc = Document(input_file)

    # Create a new document
    new_doc = Document()

    for para in original_doc.paragraphs:
        # Add a new paragraph to the new document
        new_para = new_doc.add_paragraph()

        # Loop through runs in the original paragraph to preserve the formatting
        for run in para.runs:
            # Add a run with the same text and formatting to the new paragraph
            new_run = new_para.add_run(run.text)

            # Preserve the font style
            if run.bold:
                new_run.bold = True
            if run.italic:
                new_run.italic = True
            if run.underline:
                new_run.underline = True

            # Preserve font size
            if run.font.size:
                new_run.font.size = run.font.size

            # Preserve font color
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb

            # Preserve font family (if defined)
            if run.font.name:
                new_run.font.name = run.font.name

        # Preserve the alignment of the paragraph
        if para.alignment is not None:
            new_para.alignment = para.alignment

    #Append placeholder name to destination filepath
    output_file = output_file + "TestCV.docx"
    # Save the new document
    new_doc.save(output_file)
    print(f"New document saved as {output_file}")


# Example usage
'''
input_file = "path/to/original_document.docx"
output_file = "path/to/output_document.docx"
copy_docx_with_formatting(input_file, output_file)
'''

